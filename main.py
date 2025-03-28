# BioAlign - DNA Sequence Alignment and Marking Tool
# Copyright (C) 2025  Vitus Lingott

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.

# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.

# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.


import json
import hashlib
import subprocess
import os.path
import sys
import html

from Bio import SeqIO
from Bio.Seq import Seq
from Bio.SeqRecord import SeqRecord

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import WD_STYLE_TYPE
from docx.oxml.ns import qn


def prepare_seq(seqs: dict, output_file_name: str):
    """
    Prepare DNA sequences for alignment using Clustal Omega.
    
    Converts dictionary of sequences to FASTA format and runs Clustal Omega
    to create a multiple sequence alignment file.
    
    Args:
        seqs: Dictionary with sequence names as keys and sequences as values
        output_file_name: Base name for output alignment file
    
    Raises:
        SystemExit: If Clustal Omega execution fails
    """
    if "." in output_file_name:
        input_file_name = "".join(output_file_name.split(".")[:-1]) + ".fasta"
    else:
        input_file_name = output_file_name + ".fasta"
        output_file_name = output_file_name + ".aln"
    with open(input_file_name, "w") as f:
        for key, value in seqs.items():
            SeqIO.write(SeqRecord(Seq(value.replace(" ", "").replace("\n", "").strip()), id=key), f, "fasta")
    command = (f".\\clustal-omega-1.2.2-win64\\clustalo.exe "
               f"--infile {input_file_name} "
               f"--outfile {output_file_name} "
               f"--outfmt clustal --force")
    command_return = subprocess.run(command)
    if command_return.returncode != 0:
        print(f"An error occurred while running clustal-omega.\nCommand: {command}\nReturncode: {command_return.returncode}")
        sys.exit(-1)


def prepare_formatted_seq(aln_file_name: str) -> str:
    """
    Format aligned sequences in triplet notation.
    
    Reads a Clustal alignment file and reformats it to add spaces 
    after every three nucleotides for better readability.
    
    Args:
        aln_file_name: Path to the Clustal alignment file
        
    Returns:
        Formatted alignment string with triplet notation
    """
    with open(aln_file_name, "r") as f:
        aln_lines = f.readlines()

    aln_seq = "".join(aln_lines[:3])

    i = 0
    c = ''
    while True:
        c = aln_lines[3][i]
        if c == ' ':
            break
        i += 1
    while True:
        c = aln_lines[3][i]
        if c != ' ':
            break
        i += 1

    for line in aln_lines[3:]:
        if line.strip() == "":
            aln_seq += line
            continue
        ID = line[:i]
        seq = line[i:]
        seq_n = ""
        for k, c in enumerate(seq):
            seq_n += c
            if (k + 1) % 3 == 0:
                seq_n += " "
        aln_seq += ID + seq_n.rstrip() + "\n"
    return aln_seq


def split_line(line: str):
    """
    Split an alignment line into sequence name and actual sequence.
    
    Args:
        line: A line from the alignment file
        
    Returns:
        Tuple of (name, sequence) or None if parsing fails
    """
    stripped_line = line.strip()
    parts = stripped_line.split(' ', 1)
    if len(parts) == 2 and parts[0]:
        name, sequence = parts
        return name, sequence
    return None


def mark_sequence_line(sequence_line: str, search_sequence: str, pattern_index: int = None, 
                       preliminary_start: int = None, ignore_spaces: bool = False):
    """
    Find all occurrences of a pattern in a sequence line.
    
    This function implements the core sequence search algorithm with two modes:
    1. Exact matching - spaces are significant
    2. Spaced matching - spaces are ignored during comparison
    
    The function also handles partial matches that might continue on the next line
    by tracking preliminary match state.
    
    Args:
        sequence_line: The text line containing the sequence to search
        search_sequence: The pattern to search for
        pattern_index: Index in search_sequence for continuing a previous match
        preliminary_start: Starting position of a preliminary match from previous line
        ignore_spaces: Whether to ignore spaces when matching
        
    Returns:
        Tuple of (matches, match_count, preliminary_state, preliminary_completed)
        - matches: List of (start, end) positions of matches
        - match_count: Number of matches found
        - preliminary_state: Tuple of (start, end, pattern_index) for partial match
        - preliminary_completed: Whether a preliminary match was completed
    """
    matches = []
    matches_num = 0
    len_search = len(search_sequence)
    len_seq = len(sequence_line)
    preliminary = ()  # Stores state for potential match continuation
    
    # Track if a preliminary match was completed
    # None = preliminary match in progress
    # True = preliminary match successful 
    # False = preliminary match failed or no preliminary match was attempted
    preliminary_completed = False if pattern_index is None else None

    if len_search == 0:
        return None
    current_pos = 0

    # Simple case: exact matching (spaces count as characters)
    if not ignore_spaces:
        while current_pos <= len_seq - len_search:
            found_pos = sequence_line.find(search_sequence, current_pos)
            if found_pos != -1:
                match_start, match_end = found_pos, found_pos + len_search
                matches.append((match_start, match_end))
                matches_num += 1
                current_pos = match_end  # Continue search after this match
            else:
                break
    # Complex case: ignore spaces during matching
    else:
        while current_pos < len_seq:
            # Initialize tracking variables for this match attempt
            pattern_i = 0  # Current position in pattern
            sequence_i = current_pos  # Current position in sequence
            match_start_i = -1  # Starting position of potential match
            
            # Continue a match from previous line if pattern_index is provided
            if pattern_index is not None and pattern_index != -1:
                pattern_i = pattern_index  # Resume pattern matching from this position
                if preliminary_start is not None:
                    match_start_i = preliminary_start  # Use provided start position
                pattern_index = -1  # Reset to avoid using this value again
                
            # Character-by-character comparison loop
            while pattern_i < len_search and sequence_i < len_seq:
                seq_char = sequence_line[sequence_i]
                
                # Skip spaces in the sequence
                if seq_char == ' ':
                    sequence_i += 1
                    # If we haven't started a match yet, update current position too
                    if match_start_i == -1 and pattern_i == 0:
                        current_pos = sequence_i
                    continue
                
                # Mark the start of a potential match
                if match_start_i == -1:
                    match_start_i = sequence_i
                
                # Character matches
                if seq_char == search_sequence[pattern_i]:
                    pattern_i += 1
                    sequence_i += 1
                # Character doesn't match
                else:
                    # If continuing a preliminary match, mark it as failed
                    if preliminary_completed is None:
                        preliminary_completed = False
                    break
            
            # Complete match found
            if pattern_i == len_search:
                match_end_i = sequence_i
                matches.append((match_start_i, match_end_i))
                matches_num += 1
                current_pos = match_end_i  # Continue search after this match
                
                # If this was a preliminary match continuation, mark it complete
                if preliminary_completed is None:
                    preliminary_completed = True
            # No complete match at this position
            else:
                # Track partial match at end of line for potential continuation
                if pattern_i > 0 and sequence_i == len_seq:
                    preliminary = (match_start_i, sequence_i, pattern_i)
                
                current_pos += 1  # Try next position
                
                # Store partial match data if at end of sequence
                if current_pos == len_seq and match_start_i != -1 and pattern_i > 0:
                    preliminary = (match_start_i, sequence_i, pattern_i)
                    
    # Return results
    if len(matches) != 0:
        return matches, matches_num, preliminary, preliminary_completed
    else:
        return None, 0, preliminary, preliminary_completed if preliminary_completed is not None else False


def process_lines(lines: list, search_sequence: str, sequences_num: int, 
                 ignore_spaces: bool = False) -> tuple:
    """
    Process all lines of the alignment to find pattern matches.
    
    This function handles matches that continue across multiple lines by tracking
    preliminary (partial) matches for each sequence name.
    
    Args:
        lines: All lines from the alignment file
        search_sequence: Pattern to search for
        sequences_num: Number of sequences in the alignment
        ignore_spaces: Whether to ignore spaces during matching
        
    Returns:
        Tuple of (results, matches_num):
        - results: List of matches for each line (None or list of (start, end) tuples)
        - matches_num: Total number of matches found
    """
    # Dictionary to track partial matches by sequence name
    preliminary = {}
    matches_num = 0
    results = []
    
    # Number of lines between consecutive occurrences of the same sequence
    # (includes sequences + blank/marker lines)
    lines_until_next_sequence = sequences_num + 2
    
    # Process each line of the alignment
    for i, line in enumerate(lines):
        # Skip header lines (CLUSTAL format has 3 header lines)
        if i < 3:
            results.append(None)
            continue
            
        # Extract sequence name and content
        name_sequence_line = split_line(line)
        if name_sequence_line:
            name, sequence_line = name_sequence_line
            
            # Check if we have a partial match from previous occurrence of this sequence
            if name in preliminary and preliminary.get(name) != ():
                prev_start, _, prev_pattern_i = preliminary.get(name)
                
                # Continue matching from where we left off
                # If this is the immediate next line, use preliminary_start
                matches_compound = mark_sequence_line(
                    sequence_line, 
                    search_sequence, 
                    pattern_index=prev_pattern_i, 
                    preliminary_start=prev_start if i - lines_until_next_sequence < 0 else None,
                    ignore_spaces=ignore_spaces
                )
            else:
                # Start fresh match for this sequence
                matches_compound = mark_sequence_line(sequence_line, search_sequence, ignore_spaces=ignore_spaces)
            
            # Process match results
            if matches_compound is not None:
                matches, matches_count, preliminary_value, preliminary_completed = matches_compound
                
                # Handle case where a preliminary match was completed
                # This means we need to mark the end of the previous line too
                if (preliminary_completed and 
                        name in preliminary and 
                        preliminary.get(name) != ()):
                    # Find previous line with this sequence
                    prev_line_index = i - lines_until_next_sequence
                    if prev_line_index >= 0 and prev_line_index < len(results):
                        # Initialize if needed
                        if results[prev_line_index] is None:
                            results[prev_line_index] = []
                        # Add match for previous line
                        match_start, _, _ = preliminary.get(name)
                        match_end = len(lines[prev_line_index].split(' ', 1)[1]) if ' ' in lines[prev_line_index] else 0
                        results[prev_line_index].append((match_start, match_end))
                        matches_num += 1
                
                # Update preliminary match state for this sequence
                preliminary[name] = preliminary_value
                
                # Add matches for current line
                if matches:
                    matches_num += matches_count  # Note: This could be a bug, should be matches_num not matches_count
                    results.append(matches)
                else:
                    results.append(None)
            else:
                results.append(None)
        else:
            results.append(None)
            
    return results, matches_num


def save_matches_html(
    file_path, lines, matches_results,
    page_title="BioAlign",
    font_name="Courier New, monospace",
    max_width_percent=100
):
    """
    Generate an HTML document with highlighted sequence matches.
    
    Creates an HTML file with the alignment where matched patterns
    are highlighted with yellow background.
    
    Args:
        file_path: Output HTML file path
        lines: Alignment lines to display
        matches_results: Match positions for each line
        page_title: HTML page title
        font_name: Font to use for sequence display
        max_width_percent: Maximum width of the sequence display area
        
    Returns:
        Boolean indicating success or failure
    """
    css_styles = f"""\t\t\tbody {{
                text-align: center;
                margin: 0;
                padding: 1em 0;
            }}
            pre {{
                display: inline-block;
                text-align: left;
                max-width: {max_width_percent}%;
                box-sizing: border-box;
                overflow-x: auto;
                font-family: {font_name};
                font-size: clamp(14px, 0.85vw, 18px);
            }}
            mark {{
                background-color: yellow;
                color: black;
            }}"""

    html_output_lines = [
        "<!DOCTYPE html>",
        "<html>",
        "\t<head>",
        f"\t\t<title>{html.escape(page_title)}</title>",
        "\t\t<style>",
        css_styles,
        "\t\t</style>",
        "\t</head>",
        "\t<body>",
        "\t\t<pre>"
    ]
    if len(matches_results) != 0:
        for i, line in enumerate(lines):
            cleaned_line_no_ending = line.rstrip('\n\r')
            line_matches = matches_results[i]

            if line_matches is None:
                html_output_lines.append(html.escape(cleaned_line_no_ending))
                continue

            name_sequence_pair = split_line(cleaned_line_no_ending)

            if not name_sequence_pair:
                html_output_lines.append(html.escape(cleaned_line_no_ending))
                continue

            name, sequence_line = name_sequence_pair

            line_parts = [html.escape(name) + " "]

            if not line_matches:
                line_parts.append(html.escape(sequence_line))
            else:
                last_index = 0
                for match_start, match_end in line_matches:
                    if match_start > last_index:
                        line_parts.append(html.escape(sequence_line[last_index:match_start]))

                    if match_start < match_end <= len(sequence_line):
                        line_parts.append(f"<mark>{html.escape(sequence_line[match_start:match_end])}</mark>")
                    else:
                        line_parts.append(html.escape(sequence_line[last_index:match_start]))
                        print(
                            f"Warning: Invalid match indices ({match_start}, {match_end}) for sequence length {len(sequence_line)} on line {i + 1}")

                    last_index = match_end

                if last_index < len(sequence_line):
                    line_parts.append(html.escape(sequence_line[last_index:]))

            html_output_lines.append("".join(line_parts))
    else:
        html_output_lines += lines

    html_output_lines.append("\t\t</pre>")
    html_output_lines.append("\t</body>")
    html_output_lines.append("</html>")

    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\n".join(html_output_lines))
    except Exception as e:
        print(f"Error saving HTML document '{file_path}' -> {e}")
        return False
    return True


def save_matches_word(
    file_path,
    lines,
    matches_results,
    font_name="Courier New",
    font_size=9,
    margin_inches=0.5,
    highlight_color=WD_COLOR_INDEX.YELLOW
):
    """
    Generate a Word document with highlighted sequence matches.
    
    Creates a Word document with the alignment where matched patterns
    are highlighted with the specified color.
    
    Args:
        file_path: Output DOCX file path
        lines: Alignment lines to display
        matches_results: Match positions for each line
        font_name: Font to use for sequence display
        font_size: Font size in points
        margin_inches: Document margins in inches
        highlight_color: Color to use for highlighting matches
        
    Returns:
        Boolean indicating success or failure
    """
    document = Document()

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)

    try:
        style = document.styles['Monospace']
    except KeyError:
        style = document.styles.add_style("Monospace", WD_STYLE_TYPE.PARAGRAPH) # 1 = WD_STYLE_TYPE.PARAGRAPH
        style.font.name = font_name
        style.font.size = Pt(font_size)
        rpr = style.element.get_or_add_rPr()
        rFonts = rpr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)

    if len(matches_results) != 0:
        for i, line in enumerate(lines):
            line_matches = matches_results[i]
            cleaned_line = line.rstrip('\n\r')

            paragraph = document.add_paragraph(style="Monospace")

            paragraph.paragraph_format.space_after = Pt(0)

            if line_matches is None:
                run = paragraph.add_run(cleaned_line)
                run.font.name = font_name
                run.font.size = Pt(font_size)
            else:
                name_sequence_pair = split_line(cleaned_line)

                name, sequence_line = name_sequence_pair

                run = paragraph.add_run(name + " ")
                run.font.name = font_name
                run.font.size = Pt(font_size)

                if not line_matches:
                    run = paragraph.add_run(sequence_line)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                else:
                    last_index = 0
                    for match_start, match_end in line_matches:
                        if match_start > last_index:
                            run = paragraph.add_run(sequence_line[last_index:match_start])
                            run.font.name = font_name
                            run.font.size = Pt(font_size)

                        if match_start < match_end <= len(sequence_line):
                            run = paragraph.add_run(sequence_line[match_start:match_end])
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                            if highlight_color is not None:
                                run.font.highlight_color = highlight_color
                        else:
                            print(f"Warning: Invalid match indices ({match_start}, {match_end}) for sequence length {len(sequence_line)} on line {i+1}")

                        last_index = match_end

                    if last_index < len(sequence_line):
                        run = paragraph.add_run(sequence_line[last_index:])
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
    else:
        for line in lines:
            paragraph = document.add_paragraph(style="Monospace")
            paragraph.paragraph_format.space_after = Pt(0)

            run = paragraph.add_run(line)
            run.font.name = font_name
            run.font.size = Pt(font_size)

    try:
        document.save(file_path)
    except IOError as e:
        print(f"Error saving Word document '{file_path}' -> {e}")
        return False
    except Exception as e:
        print(f"Unexpected error saving Word document '{file_path}' -> {e}")
        return False
    return True


def main():
    """
    Main program execution flow.
    
    Steps:
    1. Load sequence data from JSON
    2. Generate or reuse sequence alignment
    3. Format sequences in triplet notation
    4. Get search pattern from user
    5. Find and mark matches
    6. Generate output documents with highlighted matches
    """
    # Print license info
    print("""
    BioAlign  Copyright (C) 2025  Vitus Lingott
    This program comes with ABSOLUTELY NO WARRANTY; for details refer to the LICENSE file.
    This is free software, and you are welcome to redistribute it
    under certain conditions; refer to the LICENSE file for details.
""")

    output_path = os.path.join(os.path.dirname(__file__), "output")
    input_path = os.path.join(os.path.dirname(__file__), "input")

    os.makedirs(output_path, exist_ok=True)
    os.makedirs(input_path, exist_ok=True)

    # Load DNA sequences from JSON configuration file
    try:
        with open(os.path.join(input_path, "sequences.json"), "r") as file:
            sequences = json.load(file)
    except FileNotFoundError:
        print(f"Error: File '{os.path.join(input_path, 'sequences.json')}' not found. Refer to the README.md file for instructions.")
        return
    except json.JSONDecodeError:
        print(f"Error: Invalid JSON format in '{os.path.join(input_path, 'sequences.json')}'. Refer to the README.md file for instructions.")
        return

    # Skip regeneration of alignment if sequences haven't changed
    # This uses MD5 hash to track changes to sequence data
    path_hash = os.path.join(output_path, ".sequencehash")
    last_hash = None
    if os.path.exists(path_hash) and os.path.isfile(path_hash):
        with open(path_hash, "r", encoding="utf-8") as file:
            last_hash = file.read()
    
    # Calculate hash of current sequences
    new_hash = hashlib.md5(str(sequences).encode("utf-8"), usedforsecurity=False).hexdigest()
    
    # Generate alignment if hash changed or alignment file doesn't exist
    if last_hash != new_hash or not os.path.exists(os.path.join(output_path, "sequences.aln")) or not os.path.isfile(os.path.join(output_path, "sequences.aln")):
        with open(path_hash, "w", encoding="utf-8") as file:
            file.write(new_hash)
        print("Computing DNA sequence alignment...", end='')
        prepare_seq(sequences, os.path.join(output_path, "sequences.aln"))
        print("\rComputed DNA sequence alignment.   \n")
    else:
        print("Reusing unchanged DNA alignment file.\n")

    # Format DNA sequence into triplet notation for better readability
    text_lines = prepare_formatted_seq(os.path.join(output_path, "sequences.aln")).splitlines()

    # Define output filenames
    word_output_filename = os.path.join(output_path, "sequences.docx")
    html_output_filename = os.path.join(output_path, "sequences.html")

    # Initialize match results list
    match_results = []
    
    # Get search pattern from user
    search_word = input("Input DNA sequence to search for (e.g. \"ACC\" or \"\" to disable marking)\n > ").strip()
    print()
    
    if len(search_word) != 0:
        # Determine search mode: exact or ignoring spaces
        skip_spaces = True if "space" in input("Search mode (exact/spaced)\n > ") else False
        print()

        # Find pattern matches in the sequences
        match_results, match_num = process_lines(text_lines, search_word, len(sequences), ignore_spaces=skip_spaces)
        print(f"{match_num} matches found.\n")
    else:
        print("Entered empty search phrase.\n")
    
    # Create output documents with highlighted matches
    if save_matches_html(html_output_filename, text_lines, match_results):
        print(f"Saved HTML to '{html_output_filename}'")
    if save_matches_word(word_output_filename, text_lines, match_results):
        print(f"Saved Word document to '{word_output_filename}'")
    
    print("\nExiting...")


if __name__ == "__main__":
    main()
